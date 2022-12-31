import speech_recognition as sr
import os
from docx import Document

r=sr.Recognizer()
with sr.Microphone() as recurso:
    print('Dime algo...')
    audio = r.listen(recurso)
    try:
        texto=r.recognize_google(audio,language='es-ES')
        print('Esto es lo que has dicho: {}'.format(texto))
        
        document = Document()
        document.add_heading('Prueba', 0)
        p = document.add_paragraph('')
        document.add_paragraph(texto)
        document.save(os.environ['USERPROFILE']+'/Desktop/test.docx')
    except:
        print('Lo siento, no te entendí.')